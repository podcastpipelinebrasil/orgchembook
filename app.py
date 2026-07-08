"""
OrgChemBook — Streamlit version
================================
Reaction calculator · Stoichiometry · Green-chemistry metrics · Batch & Flow · Biocatalysis

Run locally:
    pip install -r requirements.txt
    streamlit run app.py

requirements.txt
----------------
    streamlit>=1.30
    pandas>=2.0
    requests>=2.31
    fpdf2>=2.7          # optional, only needed for PDF export
    python-docx>=1.1    # optional, only needed for Word (.docx) export

Deploy: Streamlit Community Cloud (share.streamlit.io) — just point it at app.py.

Notes
-----
* Compound lookup uses the free PubChem PUG-REST API (no API key required).
* Green-metric conventions (corrected for consistency vs. the React prototype):
    AE       = MW(product) / Σ MW(reagents) · 100
    E-factor = (total input mass − product mass) / product mass
    PMI      = total input mass / product mass          <-- uses EXPERIMENTAL product
               (falls back to theoretical yield if no experimental mass given)
    E+       = E-factor(total, incl. work-up) + E-factor(energy)
    Energy   = kWh · CI / product_mass ,  CI = 0.060 kg CO2 / kWh
"""

import io
import json
import math
import uuid
from dataclasses import dataclass, field, asdict
from typing import Optional

import pandas as pd
import requests
import streamlit as st

# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────
MW_ATOMS = {
    "C": 12.011, "H": 1.008, "N": 14.007, "O": 15.999, "S": 32.06, "P": 30.974,
    "F": 18.998, "Cl": 35.45, "Br": 79.904, "I": 126.904, "Si": 28.086, "B": 10.81,
    "Na": 22.99, "K": 39.098, "Ca": 40.078, "Mg": 24.305, "Fe": 55.845, "Cu": 63.546,
    "Zn": 65.38, "Al": 26.982, "Li": 6.941, "Ba": 137.327, "Mn": 54.938, "Co": 58.933,
    "Ni": 58.693, "Ag": 107.868, "Sn": 118.71, "Pb": 207.2, "Ti": 47.867, "Cr": 51.996,
    "Se": 78.96, "As": 74.922,
}
ROLES = ["reagent", "solvent", "catalyst", "product"]
EQP = {  # reactor -> (rated kW, load factor)
    "round-bottom flask": (0.6, 0.3),
    "jacketed reactor": (2.5, 0.5),
    "shaker": (0.8, 0.8),
}
CI = 0.06  # kg CO2 / kWh — grid carbon intensity assumption

ROLE_COLORS = {
    "reagent": "#4f46e5", "solvent": "#0891b2",
    "catalyst": "#d97706", "product": "#059669",
}

# ─────────────────────────────────────────────────────────────────────────────
# Pure chemistry helpers
# ─────────────────────────────────────────────────────────────────────────────
def parse_mw(formula: str) -> Optional[float]:
    """Compute molecular weight from a formula string (supports nested parentheses)."""
    if not formula:
        return None

    def _parse(s: str) -> Optional[float]:
        mw, i = 0.0, 0
        while i < len(s):
            c = s[i]
            if c == "(":
                depth, j = 1, i + 1
                while j < len(s) and depth > 0:
                    if s[j] == "(":
                        depth += 1
                    elif s[j] == ")":
                        depth -= 1
                    j += 1
                k, num = j, ""
                while k < len(s) and s[k].isdigit():
                    num += s[k]
                    k += 1
                inner = _parse(s[i + 1:j - 1])
                if inner is None:
                    return None
                mw += inner * (int(num) if num else 1)
                i = k
            elif c.isupper():
                sym, j = c, i + 1
                while j < len(s) and s[j].islower():
                    sym += s[j]
                    j += 1
                k, num = j, ""
                while k < len(s) and s[k].isdigit():
                    num += s[k]
                    k += 1
                if sym not in MW_ATOMS:
                    return None
                mw += MW_ATOMS[sym] * (int(num) if num else 1)
                i = k
            else:
                i += 1
        return mw

    r = _parse(formula)
    return round(r, 3) if r and r > 0 else None


def link_quantities(mw, mass, volume, mmol, density, purity, changed):
    """Recompute mass/volume/mmol keeping them consistent. Returns (mass, volume, mmol)."""
    pur = (purity or 100) / 100
    if not mw or mw <= 0:
        return mass, volume, mmol

    def to_f(x):
        try:
            return float(x)
        except (TypeError, ValueError):
            return None

    m, v, mol, d = to_f(mass), to_f(volume), to_f(mmol), to_f(density)

    if changed == "mass" and m and m > 0:
        mol = (m / mw) * 1000 * pur
        if d and d > 0:
            v = m / d
    elif changed == "volume" and v and v > 0 and d and d > 0:
        m = v * d
        mol = (m / mw) * 1000 * pur
    elif changed == "mmol" and mol and mol > 0:
        m = (mol / pur) * mw / 1000
        if d and d > 0:
            v = m / d
    elif changed in ("mw", "purity", "density"):
        if mol and mol > 0:
            m = (mol / pur) * mw / 1000
            if d and d > 0:
                v = m / d
        elif m and m > 0:
            mol = (m / mw) * 1000 * pur
            if d and d > 0:
                v = m / d

    def clean(x):
        return round(x, 5) if isinstance(x, float) else x

    return clean(m), clean(v), clean(mol)


def calc_energy(reactor, time_min, product_g, photochemical, led_power, ultrasound):
    """Energy consumption (kWh) and energy E-factor (kg CO2 / g product)."""
    try:
        t = float(time_min)
        pg = float(product_g)
    except (TypeError, ValueError):
        return None
    if t <= 0 or pg <= 0:
        return None
    t_h = t / 60
    kw, f = EQP.get(reactor, (0, 0))
    kwh = kw * f * t_h
    if photochemical and led_power:
        try:
            kwh += (float(led_power) / 1000) * t_h
        except ValueError:
            pass
    if ultrasound:
        kwh += 0.1 * 0.2 * t_h  # transducer ~100 W @ 20% duty
    return {"kwh": round(kwh, 4), "ef_energy": round((kwh * CI) / pg, 4)}


# ─────────────────────────────────────────────────────────────────────────────
# PubChem lookup (free PUG-REST, no key)
# ─────────────────────────────────────────────────────────────────────────────
PUGREST = "https://pubchem.ncbi.nlm.nih.gov/rest/pug"


@st.cache_data(show_spinner=False, ttl=86400)
def pubchem_lookup(query: str) -> dict:
    """Look up a compound by CAS or name. Returns {name, formula, mw, density, smiles, cid, error?}."""
    query = query.strip()
    if not query:
        return {"error": "empty query"}
    try:
        # Resolve to a CID (works for CAS numbers and names alike)
        r = requests.get(f"{PUGREST}/compound/name/{requests.utils.quote(query)}/cids/JSON", timeout=12)
        if r.status_code != 200:
            return {"error": f"Not found (HTTP {r.status_code})"}
        cids = r.json().get("IdentifierList", {}).get("CID", [])
        if not cids:
            return {"error": "No CID found"}
        cid = cids[0]

        # Core properties
        props = "MolecularFormula,MolecularWeight,IUPACName,CanonicalSMILES"
        p = requests.get(f"{PUGREST}/compound/cid/{cid}/property/{props}/JSON", timeout=12)
        pdata = p.json()["PropertyTable"]["Properties"][0]

        density = _pubchem_density(cid)
        common = _pubchem_common_name(cid)
        return {
            "cid": cid,
            "name": common or pdata.get("IUPACName") or query,
            "formula": pdata.get("MolecularFormula"),
            "mw": float(pdata["MolecularWeight"]) if pdata.get("MolecularWeight") else None,
            "smiles": pdata.get("CanonicalSMILES"),
            "iupac": pdata.get("IUPACName"),
            "density": density,
        }
    except requests.RequestException as e:
        return {"error": f"Network error: {e}"}
    except (KeyError, ValueError, IndexError) as e:
        return {"error": f"Parse error: {e}"}


def _pubchem_common_name(cid: int) -> Optional[str]:
    """Fetch the preferred common/synonym name for a CID (first synonym is usually the common name)."""
    try:
        r = requests.get(f"{PUGREST}/compound/cid/{cid}/synonyms/JSON", timeout=12)
        if r.status_code != 200:
            return None
        syns = r.json()["InformationList"]["Information"][0].get("Synonym", [])
        # PubChem lists the most common name first; skip pure CAS-number entries
        for s in syns:
            if not s.replace("-", "").isdigit():  # skip CAS-like "67-64-1"
                return s
        return syns[0] if syns else None
    except (requests.RequestException, KeyError, IndexError):
        return None


def pubchem_image_url(cid: int, size: int = 200) -> str:
    """Return a PubChem 2D structure image URL for a given CID."""
    return f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/cid/{cid}/PNG?image_size={size}x{size}"


# ─────────────────────────────────────────────────────────────────────────────
# Streamlit callbacks — run BEFORE the script re-executes, so they can freely
# write to widget keys (this is the reliable way to populate fields).
# ─────────────────────────────────────────────────────────────────────────────
def cb_lookup(uid, k_name, k_cas, k_form, k_mw, k_dens, k_smi, k_cid):
    """Fetch from PubChem and write results directly into the widget keys."""
    query = (st.session_state.get(k_cas) or st.session_state.get(k_name) or "").strip()
    if not query:
        return
    data = pubchem_lookup(query)
    if data.get("error"):
        st.session_state[f"_msg_{uid}"] = ("error", data["error"])
        return
    if data.get("name"):
        st.session_state[k_name] = data["name"]
    if data.get("formula"):
        st.session_state[k_form] = data["formula"]
    if data.get("mw"):
        st.session_state[k_mw] = float(data["mw"])
    if k_dens and data.get("density"):
        st.session_state[k_dens] = float(data["density"])
    if k_smi and data.get("smiles"):
        st.session_state[k_smi] = data["smiles"]
    st.session_state[k_cid] = data.get("cid")
    st.session_state[f"_msg_{uid}"] = (
        "ok", f"✔ CID {data['cid']} · {data.get('formula')} · MW {data.get('mw')}")


def cb_relink(uid, changed):
    """Recompute linked quantities (mass<->volume<->mmol) and write back to keys."""
    mw = st.session_state.get(f"mw_{uid}")
    if not mw:
        return
    mass = st.session_state.get(f"mass_{uid}") or None
    vol = st.session_state.get(f"vol_{uid}") or None
    mmol = st.session_state.get(f"mmol_{uid}") or None
    dens = st.session_state.get(f"dens_{uid}") or None
    pur = st.session_state.get(f"pur_{uid}", 100.0)
    lm, lv, lmol = link_quantities(mw, mass, vol, mmol, dens, pur, changed)
    st.session_state[f"mass_{uid}"] = float(lm) if lm else 0.0
    st.session_state[f"vol_{uid}"] = float(lv) if lv else 0.0
    st.session_state[f"mmol_{uid}"] = float(lmol) if lmol else 0.0


def _pubchem_density(cid: int) -> Optional[float]:
    """Try to extract a numeric density (g/mL) from the PubChem PUG-View record."""
    try:
        url = f"https://pubchem.ncbi.nlm.nih.gov/rest/pug_view/data/compound/{cid}/JSON?heading=Density"
        r = requests.get(url, timeout=12)
        if r.status_code != 200:
            return None
        txt = json.dumps(r.json())
        import re
        # grab first number that looks like a density between 0.3 and 3.5 g/cm3
        for m in re.finditer(r"(\d+\.?\d*)\s*(?:g/cm3|g/mL|g/ml)", txt):
            val = float(m.group(1))
            if 0.3 <= val <= 3.5:
                return round(val, 3)
    except (requests.RequestException, ValueError, KeyError):
        return None
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Data model
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class Compound:
    name: str = ""
    cas: str = ""
    formula: str = ""
    smiles: str = ""
    mw: Optional[float] = None
    mass: Optional[float] = None
    volume: Optional[float] = None
    mmol: Optional[float] = None
    density: Optional[float] = None
    purity: float = 100.0
    role: str = "reagent"
    cid: Optional[int] = None  # PubChem CID (for structure image)
    # unique, stable id used for widget keys (never changes for the life of the object)
    uid: str = field(default_factory=lambda: uuid.uuid4().hex[:8])
    # biocatalyst
    cat_type: str = "chemical"
    enzyme_name: str = ""
    enzyme_code: str = ""
    enzyme_form: str = "solid"
    enzyme_amount: Optional[float] = None
    enzyme_activity: Optional[float] = None
    is_limiting: bool = False


def default_reagents():
    return [Compound(role="reagent"), Compound(role="product")]


# ─────────────────────────────────────────────────────────────────────────────
# Metric status helpers (for color coding)
# ─────────────────────────────────────────────────────────────────────────────
def ae_status(v):   return "🟢" if v >= 80 else ("🟡" if v >= 50 else "🔴")
def ef_status(v):   return "🟢" if v < 5 else ("🟡" if v < 25 else "🔴")
def pmi_status(v):  return "🟢" if v < 10 else ("🟡" if v < 50 else "🔴")


# ─────────────────────────────────────────────────────────────────────────────
# Core calculation
# ─────────────────────────────────────────────────────────────────────────────
def compute_batch(reagents, workup, cond, exp_yield):
    non_p = [r for r in reagents if r.role != "product"]
    prods = [r for r in reagents if r.role == "product"]

    reagent_pool = [r for r in non_p if r.role == "reagent" and r.mmol]
    if not reagent_pool:
        return {"error": "Add at least one reagent with amount and MW."}

    limiting = next((r for r in reagent_pool if r.is_limiting), None)
    if limiting is None:
        limiting = min(reagent_pool, key=lambda r: r.mmol)

    # equivalents
    rows = []
    for r in non_p:
        equiv = (r.mmol / limiting.mmol) if (r.mmol and limiting.mmol) else None
        rows.append((r, equiv))

    # product MW
    prod_mw = next((p.mw or parse_mw(p.formula) for p in prods if (p.mw or parse_mw(p.formula))), None)
    reagent_mws = [(r.mw or parse_mw(r.formula)) for r in reagents
                   if r.role == "reagent" and (r.mw or parse_mw(r.formula))]
    sum_mw = sum(reagent_mws)

    ae = round((prod_mw / sum_mw) * 100, 1) if (prod_mw and sum_mw) else None
    ty = (limiting.mmol * prod_mw / 1000) if prod_mw else None          # g
    tym = limiting.mmol                                                  # mmol

    reaction_input = sum((r.mmol * r.mw / 1000) for r in non_p if r.mmol and r.mw)
    wu_mass = sum(w["amount"] for w in workup if w.get("amount"))
    total_input = reaction_input + wu_mass

    py = ef = ef_wu = pmi = pmi_wu = None
    ey = float(exp_yield) if exp_yield else None

    # ── Corrected metric conventions ─────────────────────────────────────────
    # PMI uses the EXPERIMENTAL product mass when available (consistent w/ E-factor),
    # otherwise falls back to theoretical yield.
    product_basis = ey if ey else ty

    if ey and ty:
        py = round((ey / ty) * 100, 1)
    if ey:
        ef = round((reaction_input - ey) / ey, 2)
        ef_wu = round((total_input - ey) / ey, 2)
    if product_basis:
        pmi = round(reaction_input / product_basis, 2)
        pmi_wu = round(total_input / product_basis, 2)

    energy = calc_energy(cond["reactor"], cond["time"], exp_yield,
                         cond["photochemical"], cond["led_power"], cond["ultrasound"])
    base_ef = ef_wu if ef_wu is not None else ef
    ef_plus = round(base_ef + energy["ef_energy"], 4) if (energy and base_ef is not None) else None

    biocats = [r for r in reagents
               if r.role == "catalyst" and r.cat_type == "biological" and r.enzyme_name]

    return {
        "mode": "batch", "rows": rows, "limiting": limiting,
        "ae": ae, "ty": ty, "tym": tym, "py": py,
        "ef": ef, "ef_wu": ef_wu, "pmi": pmi, "pmi_wu": pmi_wu,
        "reaction_input": reaction_input, "wu_mass": wu_mass, "total_input": total_input,
        "energy": energy, "ef_plus": ef_plus, "biocats": biocats,
        "product_basis_label": "experimental" if ey else "theoretical",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Export helpers
# ─────────────────────────────────────────────────────────────────────────────
def build_summary_df(reagents):
    data = []
    for r in reagents:
        data.append({
            "Name": r.enzyme_name if (r.cat_type == "biological") else r.name,
            "CAS": r.cas, "Role": r.role, "Formula": r.formula,
            "MW": r.mw, "Mass (g)": r.mass, "Volume (mL)": r.volume,
            "mmol": r.mmol, "Density": r.density, "Purity (%)": r.purity,
        })
    return pd.DataFrame(data)


def build_pdf(meta, res, df):
    try:
        from fpdf import FPDF
    except ImportError:
        return None
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 15)
    pdf.cell(0, 10, "OrgChemBook - Reaction Report", ln=True)
    pdf.set_font("Helvetica", "", 10)
    for k, v in meta.items():
        if v:
            pdf.cell(0, 6, f"{k}: {v}", ln=True)
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Green-chemistry metrics", ln=True)
    pdf.set_font("Helvetica", "", 10)
    metrics = [
        ("Atom economy (%)", res.get("ae")),
        ("E-factor (reaction)", res.get("ef")),
        ("E-factor (+ work-up)", res.get("ef_wu")),
        ("PMI", res.get("pmi")),
        ("E+ factor", res.get("ef_plus")),
        ("Theoretical yield (g)", round(res["ty"], 4) if res.get("ty") else None),
        ("% Yield", res.get("py")),
    ]
    for label, val in metrics:
        if val is not None:
            pdf.cell(0, 6, f"  {label}: {val}", ln=True)
    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Compounds", ln=True)
    pdf.set_font("Helvetica", "", 8)
    for _, row in df.iterrows():
        line = f"  {row['Role']}: {row['Name'] or '-'} | MW {row['MW'] or '-'} | {row['mmol'] or '-'} mmol"
        pdf.cell(0, 5, line, ln=True)
    return bytes(pdf.output())


def build_docx_flow(meta, res):
    """Editable Word (.docx) report for FLOW mode. Returns bytes or None if python-docx missing."""
    try:
        from docx import Document
    except ImportError:
        return None

    doc = Document()
    doc.add_heading("OrgChemBook — Flow Reaction Report", level=0)

    for k, v in meta.items():
        if v:
            p = doc.add_paragraph()
            p.add_run(f"{k}: ").bold = True
            p.add_run(str(v))

    # Pumps table
    doc.add_heading("Flow setup — pumps", level=1)
    cols = ["Pump", "Role", "Compound", "CAS", "MW", "mmol", "Limiting"]
    pt = doc.add_table(rows=1, cols=len(cols))
    pt.style = "Light Grid Accent 1"
    for j, c in enumerate(cols):
        pt.rows[0].cells[j].paragraphs[0].add_run(c).bold = True
    for i, p in enumerate(res["pumps"]):
        name = p.enzyme_name if (p.role == "catalyst" and p.cat_type == "biological") else p.name
        vals = [f"P{str(i+1).zfill(2)}", p.role, name or "—", p.cas or "—",
                str(p.mw or "—"), str(p.mmol or "—"), "★" if p.is_limiting else ""]
        cells = pt.add_row().cells
        for j, v in enumerate(vals):
            cells[j].text = v

    # Product
    fp = res["flow_product"]
    doc.add_heading("Product", level=1)
    doc.add_paragraph(f"Name: {fp.name or '—'} · Formula: {fp.formula or '—'} · MW: {fp.mw or '—'}")
    if res.get("exp_yield"):
        doc.add_paragraph(f"Experimental yield: {res['exp_yield']} g")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx(meta, res, df):
    """Editable Word (.docx) report for BATCH mode. Returns bytes or None if python-docx missing."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()

    # Title
    title = doc.add_heading("OrgChemBook — Reaction Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    for k, v in meta.items():
        if v:
            p = doc.add_paragraph()
            run = p.add_run(f"{k}: ")
            run.bold = True
            p.add_run(str(v))

    # Green-chemistry metrics
    doc.add_heading("Green-chemistry metrics", level=1)
    metrics = [
        ("Atom economy (%)", res.get("ae")),
        ("E-factor (reaction)", res.get("ef")),
        ("E-factor (+ work-up)", res.get("ef_wu")),
        ("PMI", res.get("pmi")),
        ("PMI (+ work-up)", res.get("pmi_wu")),
        ("E⁺ factor", res.get("ef_plus")),
        ("Theoretical yield (g)", round(res["ty"], 4) if res.get("ty") else None),
        ("% Yield", res.get("py")),
    ]
    mt = doc.add_table(rows=1, cols=2)
    mt.style = "Light Grid Accent 1"
    hdr = mt.rows[0].cells
    hdr[0].paragraphs[0].add_run("Metric").bold = True
    hdr[1].paragraphs[0].add_run("Value").bold = True
    for label, val in metrics:
        if val is not None:
            row = mt.add_row().cells
            row[0].text = label
            row[1].text = str(val)

    energy = res.get("energy")
    if energy:
        doc.add_paragraph(
            f"Energy: {energy['kwh']} kWh · E-factor(energy) = "
            f"{energy['ef_energy']} kg CO₂/g (CI = {CI})"
        )

    # Compounds table
    doc.add_heading("Compounds", level=1)
    ct = doc.add_table(rows=1, cols=len(df.columns))
    ct.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        ct.rows[0].cells[j].paragraphs[0].add_run(str(col)).bold = True
    for _, row in df.iterrows():
        cells = ct.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            cells[j].text = "" if pd.isna(val) else str(val)

    # Biocatalyst table
    biocats = res.get("biocats") or []
    if biocats:
        doc.add_heading("Biocatalyst summary", level=1)
        lm = res["limiting"].mmol
        bcols = ["Enzyme", "EC", "Form", "Amount", "Loading (mg/mmol)", "Total (U)"]
        bt = doc.add_table(rows=1, cols=len(bcols))
        bt.style = "Light Grid Accent 1"
        for j, c in enumerate(bcols):
            bt.rows[0].cells[j].paragraphs[0].add_run(c).bold = True
        for x in biocats:
            amt = x.enzyme_amount
            loading = round(amt * 1000 / lm, 1) if (amt and lm) else "—"
            total_u = round(amt * 1000 * x.enzyme_activity, 0) if (amt and x.enzyme_activity) else "—"
            unit = "mL" if x.enzyme_form == "liquid" else "g"
            vals = [x.enzyme_name or "—", x.enzyme_code or "—", x.enzyme_form,
                    f"{amt} {unit}" if amt else "—", str(loading), str(total_u)]
            cells = bt.add_row().cells
            for j, v in enumerate(vals):
                cells[j].text = v

    # Reaction conditions
    cond = res.get("cond")
    if cond:
        doc.add_heading("Reaction conditions", level=1)
        cond_lines = [
            ("Reactor", cond.get("reactor")),
            ("Agitation", cond.get("agitation")),
            ("Time (min)", cond.get("time")),
            ("Temperature (°C)", cond.get("temperature")),
            ("Photochemical", "Yes" if cond.get("photochemical") else "No"),
            ("Ultrasound", "Yes" if cond.get("ultrasound") else "No"),
        ]
        for label, val in cond_lines:
            if val not in (None, "", 0):
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(val))
        if cond.get("procedure"):
            p = doc.add_paragraph()
            p.add_run("Procedure notes: ").bold = True
            p.add_run(cond["procedure"])

    # Work-up
    workup = [w for w in st.session_state.workup if w.get("name") or w.get("amount")]
    if workup:
        doc.add_heading("Work-up", level=1)
        wt = doc.add_table(rows=1, cols=3)
        wt.style = "Light Grid Accent 1"
        for j, c in enumerate(["Material", "Amount", "Unit"]):
            wt.rows[0].cells[j].paragraphs[0].add_run(c).bold = True
        for w in workup:
            cells = wt.add_row().cells
            cells[0].text = w.get("name") or "—"
            cells[1].text = str(w.get("amount") or "—")
            cells[2].text = w.get("unit", "")

    # Method notes
    doc.add_heading("Notes", level=1)
    doc.add_paragraph(
        "AE = MW(product) / Σ MW(reagents) · E-factor = waste / product · "
        "PMI = total input mass / product · E⁺ = E-factor(total) + E-factor(energy). "
        f"PMI computed on {res.get('product_basis_label', 'theoretical')} product mass."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Streamlit UI
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="OrgChemBook", page_icon="⚗️", layout="wide")

# ── Session state init ───────────────────────────────────────────────────────
# Bump this when the Compound dataclass gains new fields, to flush stale objects
# left in session_state from an older version of the app.
SCHEMA_VERSION = 5
if st.session_state.get("_schema") != SCHEMA_VERSION:
    st.session_state.clear()
    st.session_state["_schema"] = SCHEMA_VERSION

if "reagents" not in st.session_state:
    st.session_state.reagents = default_reagents()
if "pumps" not in st.session_state:
    st.session_state.pumps = [Compound(role="reagent")]
if "flow_product" not in st.session_state:
    st.session_state.flow_product = Compound(role="product")
if "workup" not in st.session_state:
    st.session_state.workup = [{"name": "", "amount": None, "unit": "mL"}]
if "results" not in st.session_state:
    st.session_state.results = None

# ── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("⚗️ OrgChemBook")
    st.caption("Stoichiometry · Green metrics · Batch & Flow · Biocatalysis")
    st.divider()
    mode = st.radio("Reaction mode", ["Batch", "Continuous flow"], horizontal=False)
    mode = "batch" if mode == "Batch" else "flow"
    st.divider()
    st.subheader("Metadata")
    operator = st.text_input("Operator", key="operator")
    institution = st.text_input("Institution / company", key="institution")
    reaction_name = st.text_input("Reaction name", key="reaction_name")
    st.divider()
    st.caption("💡 Compound lookup uses the free PubChem PUG-REST API — no key needed.")

tab_setup, tab_results = st.tabs(["⚗️ Reaction setup", "📊 Results & metrics"])

# ═════════════════════════════════════════════════════════════════════════════
# TAB: SETUP
# ═════════════════════════════════════════════════════════════════════════════
with tab_setup:
    st.subheader("Reaction scheme")
    scheme = st.file_uploader("Upload reaction scheme (optional)", type=["png", "jpg", "jpeg", "svg"])
    if scheme:
        st.image(scheme, use_container_width=True)

    # ── BATCH MODE ───────────────────────────────────────────────────────────
    if mode == "batch":
        st.markdown("#### Compounds")
        st.caption("Order: reagents → catalysts → solvents · product last")

        # add-compound buttons
        cols = st.columns(4)
        for i, role in enumerate(ROLES):
            if cols[i].button(f"➕ Add {role}", key=f"add_{role}", use_container_width=True):
                new = Compound(role=role)
                if role == "product":
                    st.session_state.reagents.append(new)
                else:
                    idx = next((k for k, r in enumerate(st.session_state.reagents)
                                if r.role == "product"), len(st.session_state.reagents))
                    st.session_state.reagents.insert(idx, new)
                st.rerun()

        # render sorted: non-products first, product last
        ordered = [r for r in st.session_state.reagents if r.role != "product"] + \
                  [r for r in st.session_state.reagents if r.role == "product"]

        for r in ordered:
            real_idx = st.session_state.reagents.index(r)
            uid = r.uid
            color = ROLE_COLORS[r.role]

            with st.container(border=True):
                head = st.columns([2, 3, 1])
                head[0].markdown(
                    f"<span style='background:{color};color:#fff;padding:2px 12px;"
                    f"border-radius:12px;font-size:12px;font-weight:700;text-transform:uppercase'>"
                    f"{r.role}</span>", unsafe_allow_html=True)
                r.role = head[1].selectbox("Role", ROLES, index=ROLES.index(r.role),
                                           key=f"role_{real_idx}", label_visibility="collapsed")
                if len(st.session_state.reagents) > 1:
                    if head[2].button("🗑️", key=f"del_{real_idx}", use_container_width=True):
                        st.session_state.reagents.pop(real_idx)
                        st.rerun()

                is_bio = r.role == "catalyst"
                if is_bio:
                    r.cat_type = st.radio("Catalyst type", ["chemical", "biological"],
                                          index=0 if r.cat_type == "chemical" else 1,
                                          horizontal=True, key=f"ctype_{real_idx}")

                if is_bio and r.cat_type == "biological":
                    b = st.columns(4)
                    r.enzyme_name = b[0].text_input("Enzyme", r.enzyme_name, key=f"en_{real_idx}")
                    r.enzyme_code = b[1].text_input("EC / ID", r.enzyme_code, key=f"ec_{real_idx}")
                    r.enzyme_form = b[2].selectbox("Form", ["solid", "liquid"],
                                                   index=0 if r.enzyme_form == "solid" else 1,
                                                   key=f"ef_{real_idx}")
                    unit = "mL" if r.enzyme_form == "liquid" else "g"
                    r.enzyme_amount = b[3].number_input(f"Amount ({unit})", value=r.enzyme_amount or 0.0,
                                                        min_value=0.0, key=f"ea_{real_idx}") or None
                    r.enzyme_activity = st.number_input("Specific activity (U/mg)",
                                                        value=r.enzyme_activity or 0.0, min_value=0.0,
                                                        key=f"eact_{real_idx}") or None
                else:
                    c1 = st.columns([2, 2, 1])
                    r.name = c1[0].text_input("Compound name", r.name, key=f"name_{real_idx}")
                    r.cas = c1[1].text_input("CAS / name", r.cas, key=f"cas_{real_idx}")
                    if c1[2].button("🔍 Lookup", key=f"lk_{real_idx}", use_container_width=True):
                        q = r.cas or r.name
                        with st.spinner(f"Searching PubChem for '{q}'…"):
                            data = pubchem_lookup(q)
                        if data.get("error"):
                            st.warning(f"Lookup failed: {data['error']}")
                        else:
                            # stash result; it is applied at the top of the next run,
                            # before any widget with these keys is instantiated.
                            st.session_state[f"_pending_lookup_{real_idx}"] = data
                            st.session_state[f"_lookup_ok_{real_idx}"] = (
                                f"✔ CID {data['cid']} · {data.get('formula')} · MW {data.get('mw')}")
                            st.rerun()

                    # show success message that survived the rerun
                    if st.session_state.get(f"_lookup_ok_{real_idx}"):
                        st.success(st.session_state.pop(f"_lookup_ok_{real_idx}"))

                    c2 = st.columns(2)
                    r.formula = c2[0].text_input("Molecular formula", r.formula, key=f"form_{real_idx}")
                    parsed = parse_mw(r.formula)
                    mw_default = float(r.mw) if r.mw else (parsed or 0.0)
                    mw_val = st.number_input("MW (g/mol)", value=mw_default,
                                             min_value=0.0, key=f"mw_{real_idx}")
                    r.mw = mw_val or parsed or None
                    r.smiles = c2[1].text_input("SMILES", r.smiles, key=f"smi_{real_idx}")

                    # Structure preview (PubChem PNG)
                    if getattr(r, "cid", None):
                        st.image(pubchem_image_url(r.cid), width=160,
                                 caption=f"CID {r.cid}")

                    # ── Product: live theoretical yield from limiting reagent ──
                    if r.role == "product":
                        _reagent_pool = [x for x in st.session_state.reagents
                                         if x.role == "reagent" and x.mmol]
                        _lim = next((x for x in _reagent_pool if x.is_limiting), None)
                        if _lim is None and _reagent_pool:
                            _lim = min(_reagent_pool, key=lambda x: x.mmol)
                        if _lim and r.mw:
                            ty_g = _lim.mmol * r.mw / 1000       # theoretical mass (g)
                            st.info(
                                f"**Expected mass (theoretical yield):** {ty_g:.4f} g "
                                f"({_lim.mmol:.4f} mmol)  \n"
                                f"Based on limiting reagent "
                                f"**{_lim.name or _lim.formula or 'reagent'}** "
                                f"× MW(product) {r.mw:g}"
                            )
                        elif r.mw and not _lim:
                            st.caption("💡 Set a limiting reagent with amount to compute expected mass.")
                        elif not r.mw:
                            st.caption("💡 Enter the product MW to compute expected mass.")

                    if r.role != "product":
                        st.caption("Quantities — auto-linked (enter MW first)")
                        q = st.columns(5)
                        new_mass = q[0].number_input("Mass (g)", value=r.mass or 0.0, min_value=0.0,
                                                     format="%.4f", key=f"mass_{real_idx}") or None
                        new_vol = q[1].number_input("Volume (mL)", value=r.volume or 0.0, min_value=0.0,
                                                    format="%.4f", key=f"vol_{real_idx}") or None
                        new_mmol = q[2].number_input("mmol", value=r.mmol or 0.0, min_value=0.0,
                                                     format="%.4f", key=f"mmol_{real_idx}") or None
                        r.density = q[3].number_input("Density (g/mL)", value=r.density or 0.0,
                                                      min_value=0.0, key=f"dens_{real_idx}") or None
                        r.purity = q[4].number_input("Purity (%)", value=r.purity, min_value=0.0,
                                                     max_value=100.0, key=f"pur_{real_idx}")

                        # detect which field the user changed, relink, write back to keys, rerun
                        changed = None
                        if new_mass != r.mass:
                            changed = "mass"
                        elif new_vol != r.volume:
                            changed = "volume"
                        elif new_mmol != r.mmol:
                            changed = "mmol"

                        if changed and r.mw:
                            lm, lv, lmol = link_quantities(
                                r.mw, new_mass, new_vol, new_mmol, r.density, r.purity, changed)
                            r.mass, r.volume, r.mmol = lm, lv, lmol
                            # stash and rerun; applied at top of next run before widgets
                            st.session_state[f"_pending_qty_{real_idx}"] = {
                                "mass": lm, "volume": lv, "mmol": lmol}
                            st.rerun()
                        else:
                            r.mass, r.volume, r.mmol = new_mass, new_vol, new_mmol

                        if r.role == "reagent":
                            r.is_limiting = st.checkbox("★ Limiting reagent", value=r.is_limiting,
                                                        key=f"lim_{real_idx}")

        # ── Batch conditions ─────────────────────────────────────────────────
        st.markdown("#### Reaction conditions (batch)")
        with st.container(border=True):
            cc = st.columns(4)
            reactor = cc[0].selectbox("Reactor", list(EQP.keys()))
            agitation = cc[1].selectbox("Agitation", ["magnetic", "mechanical", "orbital"])
            time_min = cc[2].number_input("Time (min)", value=0.0, min_value=0.0)
            temperature = cc[3].number_input("Temperature (°C)", value=25.0)
            cc2 = st.columns(4)
            photochemical = cc2[0].checkbox("Photochemical")
            wavelength = cc2[1].number_input("Wavelength (nm)", value=0.0, min_value=0.0) if photochemical else None
            led_power = cc2[2].number_input("LED power (W)", value=0.0, min_value=0.0) if photochemical else None
            ultrasound = cc2[3].checkbox("Ultrasound")
            procedure = st.text_area("Procedure notes", height=80)
        cond = {
            "reactor": reactor, "agitation": agitation, "time": time_min,
            "temperature": temperature, "photochemical": photochemical,
            "wavelength": wavelength, "led_power": led_power,
            "ultrasound": ultrasound, "procedure": procedure,
        }

    # ── FLOW MODE ────────────────────────────────────────────────────────────
    else:
        st.info("⟶ **Continuous flow** — one card per inlet stream · product last")
        st.markdown("#### Pumps")

        if st.button("➕ Add pump"):
            st.session_state.pumps.append(Compound(role="reagent"))
            st.rerun()

        for idx, p in enumerate(st.session_state.pumps):
            # Apply pending lookup before widgets
            pend = f"_pending_plookup_{idx}"
            if st.session_state.get(pend):
                d = st.session_state.pop(pend)
                if d.get("name"):
                    p.name = d["name"]
                if d.get("formula"):
                    p.formula = d["formula"]
                if d.get("mw"):
                    p.mw = d["mw"]
                if d.get("density"):
                    p.density = d["density"]
                if d.get("smiles"):
                    p.smiles = d["smiles"]
                p.cid = d.get("cid")
                for wk in ("pname", "pcas", "pform", "pmw", "pdens"):
                    st.session_state.pop(f"{wk}_{idx}", None)

            with st.container(border=True):
                head = st.columns([1, 2, 2, 1])
                head[0].markdown(f"**P{str(idx+1).zfill(2)}**")
                p.role = head[1].selectbox("Role", ["reagent", "catalyst"],
                                           index=0 if p.role == "reagent" else 1,
                                           key=f"prole_{idx}", label_visibility="collapsed")
                if p.role == "catalyst":
                    p.cat_type = head[2].selectbox("Type", ["chemical", "biological"],
                                                   index=0 if p.cat_type == "chemical" else 1,
                                                   key=f"pctype_{idx}", label_visibility="collapsed")
                if len(st.session_state.pumps) > 1:
                    if head[3].button("🗑️", key=f"pdel_{idx}", use_container_width=True):
                        st.session_state.pumps.pop(idx)
                        st.rerun()

                if p.role == "catalyst" and p.cat_type == "biological":
                    b = st.columns(4)
                    p.enzyme_name = b[0].text_input("Enzyme", p.enzyme_name, key=f"pen_{idx}")
                    p.enzyme_code = b[1].text_input("EC / ID", p.enzyme_code, key=f"pec_{idx}")
                    p.enzyme_amount = b[2].number_input("Amount (mL)", value=p.enzyme_amount or 0.0,
                                                        min_value=0.0, key=f"pea_{idx}") or None
                    p.enzyme_activity = b[3].number_input("Activity (U/mg)", value=p.enzyme_activity or 0.0,
                                                          min_value=0.0, key=f"peact_{idx}") or None
                else:
                    c1 = st.columns([2, 2, 1])
                    p.name = c1[0].text_input("Compound", p.name, key=f"pname_{idx}")
                    p.cas = c1[1].text_input("CAS / name", p.cas, key=f"pcas_{idx}")
                    if c1[2].button("🔍 Lookup", key=f"plk_{idx}", use_container_width=True):
                        with st.spinner("Searching PubChem…"):
                            data = pubchem_lookup(p.cas or p.name)
                        if data.get("error"):
                            st.warning(f"Lookup failed: {data['error']}")
                        else:
                            st.session_state[f"_pending_plookup_{idx}"] = data
                            st.session_state[f"_plookup_ok_{idx}"] = (
                                f"✔ {data.get('formula')} · MW {data.get('mw')}")
                            st.rerun()

                    if st.session_state.get(f"_plookup_ok_{idx}"):
                        st.success(st.session_state.pop(f"_plookup_ok_{idx}"))

                    c2 = st.columns(2)
                    p.formula = c2[0].text_input("Formula", p.formula, key=f"pform_{idx}")
                    parsed = parse_mw(p.formula)
                    p.mw = st.number_input("MW (g/mol)", value=float(p.mw) if p.mw else (parsed or 0.0),
                                           min_value=0.0, key=f"pmw_{idx}") or parsed or None
                    if getattr(p, "cid", None):
                        st.image(pubchem_image_url(p.cid), width=140, caption=f"CID {p.cid}")
                    q = st.columns(4)
                    p.mass = q[0].number_input("Mass (g)", value=p.mass or 0.0, min_value=0.0, key=f"pmass_{idx}") or None
                    p.mmol = q[1].number_input("mmol", value=p.mmol or 0.0, min_value=0.0, key=f"pmmol_{idx}") or None
                    p.density = q[2].number_input("Density", value=p.density or 0.0, min_value=0.0, key=f"pdens_{idx}") or None
                    p.is_limiting = q[3].checkbox("★ Limiting", value=p.is_limiting, key=f"plim_{idx}")

        # flow product
        st.markdown("#### Product")
        fp = st.session_state.flow_product
        # apply pending lookup before widgets
        if st.session_state.get("_pending_fp"):
            d = st.session_state.pop("_pending_fp")
            if d.get("name"):
                fp.name = d["name"]
            if d.get("formula"):
                fp.formula = d["formula"]
            if d.get("mw"):
                fp.mw = d["mw"]
            fp.smiles = d.get("smiles") or fp.smiles
            fp.cid = d.get("cid")
            for wk in ("fp_name", "fp_cas", "fp_form", "fp_mw"):
                st.session_state.pop(wk, None)
        with st.container(border=True):
            c = st.columns([2, 2, 1])
            fp.name = c[0].text_input("Product name", fp.name, key="fp_name")
            fp.cas = c[1].text_input("CAS / name", fp.cas, key="fp_cas")
            if c[2].button("🔍 Lookup", key="fp_lk", use_container_width=True):
                with st.spinner("Searching PubChem…"):
                    data = pubchem_lookup(fp.cas or fp.name)
                if data.get("error"):
                    st.warning(f"Lookup failed: {data['error']}")
                else:
                    st.session_state["_pending_fp"] = data
                    st.session_state["_fp_ok"] = f"✔ {data.get('formula')} · MW {data.get('mw')}"
                    st.rerun()
            if st.session_state.get("_fp_ok"):
                st.success(st.session_state.pop("_fp_ok"))

            c2 = st.columns(2)
            fp.formula = c2[0].text_input("Formula", fp.formula, key="fp_form")
            fp.mw = c2[1].number_input("MW (g/mol)",
                                       value=float(fp.mw) if fp.mw else (parse_mw(fp.formula) or 0.0),
                                       min_value=0.0, key="fp_mw") or None
            if getattr(fp, "cid", None):
                st.image(pubchem_image_url(fp.cid), width=140, caption=f"CID {fp.cid}")

            # Expected mass from limiting pump
            _pump_pool = [p for p in st.session_state.pumps if p.role == "reagent" and p.mmol]
            _limp = next((p for p in _pump_pool if p.is_limiting), None)
            if _limp is None and _pump_pool:
                _limp = min(_pump_pool, key=lambda p: p.mmol)
            if _limp and fp.mw:
                ty_g = _limp.mmol * fp.mw / 1000
                st.info(f"**Expected mass (theoretical yield):** {ty_g:.4f} g "
                        f"({_limp.mmol:.4f} mmol) — from limiting pump "
                        f"**{_limp.name or _limp.formula or 'reagent'}**")

        st.markdown("#### Reaction conditions (flow)")
        with st.container(border=True):
            fcols = st.columns(4)
            mixer = fcols[0].selectbox("Mixer", ["T-mixer", "Y-mixer", "other"])
            reactor_type = fcols[1].selectbox("Reactor", ["coil", "plate", "microchip", "packed bed"])
            reactor_volume = fcols[2].number_input("Reactor volume (mL)", value=0.0, min_value=0.0)
            residence = fcols[3].number_input("Residence time (min)", value=0.0, min_value=0.0)
            fcols2 = st.columns(4)
            pressure = fcols2[0].number_input("Pressure (bar)", value=1.0, min_value=0.0)
            f_temp = fcols2[1].number_input("Temperature (°C)", value=25.0, key="flow_temp")
            f_photo = fcols2[2].checkbox("Photochemical", key="flow_photo")
            f_us = fcols2[3].checkbox("Ultrasound", key="flow_us")
            electro = st.checkbox("Electrochemistry")
            flow_proc = st.text_area("Procedure notes", height=80, key="flow_proc")

    # ── Work-up (shared) ─────────────────────────────────────────────────────
    st.markdown("#### Work-up")
    st.caption("Materials for isolation — included in waste and PMI")
    if st.button("➕ Add work-up entry"):
        st.session_state.workup.append({"name": "", "amount": None, "unit": "mL"})
        st.rerun()
    for widx, w in enumerate(st.session_state.workup):
        wc = st.columns([3, 1, 1, 0.5])
        w["name"] = wc[0].text_input("Material", w["name"], key=f"wname_{widx}", label_visibility="collapsed")
        w["amount"] = wc[1].number_input("Amount", value=w["amount"] or 0.0, min_value=0.0,
                                         key=f"wamt_{widx}", label_visibility="collapsed") or None
        w["unit"] = wc[2].selectbox("Unit", ["mL", "g"], index=0 if w["unit"] == "mL" else 1,
                                    key=f"wunit_{widx}", label_visibility="collapsed")
        if wc[3].button("✕", key=f"wdel_{widx}"):
            st.session_state.workup.pop(widx)
            st.rerun()

    # ── Experimental yield & Calculate ───────────────────────────────────────
    st.divider()
    exp_yield = st.number_input("Experimental yield (g) — optional, enables % yield & E-factor",
                                value=0.0, min_value=0.0, format="%.4f")
    exp_yield = exp_yield or None

    if st.button("Calculate →", type="primary", use_container_width=True):
        if mode == "batch":
            res = compute_batch(st.session_state.reagents, st.session_state.workup, cond, exp_yield)
            if res.get("error"):
                st.error(res["error"])
            else:
                res["cond"] = cond
                res["meta"] = {"Operator": operator, "Institution": institution,
                               "Reaction": reaction_name}
                st.session_state.results = res
                st.success("✔ Calculated — see the **Results & metrics** tab.")
        else:
            st.session_state.results = {
                "mode": "flow", "pumps": list(st.session_state.pumps),
                "flow_product": st.session_state.flow_product,
                "meta": {"Operator": operator, "Institution": institution, "Reaction": reaction_name},
                "exp_yield": exp_yield,
            }
            st.success("✔ Saved — see the **Results & metrics** tab.")

# ═════════════════════════════════════════════════════════════════════════════
# TAB: RESULTS
# ═════════════════════════════════════════════════════════════════════════════
with tab_results:
    res = st.session_state.results
    if not res:
        st.info("Set up your reaction and click **Calculate →**.")
    elif res["mode"] == "batch":
        st.subheader("Green-chemistry metrics")
        m = st.columns(4)
        if res["ae"] is not None:
            m[0].metric(f"{ae_status(res['ae'])} Atom economy", f"{res['ae']} %")
        if res["ef"] is not None:
            m[1].metric(f"{ef_status(res['ef'])} E-factor (rxn)", res["ef"])
        if res["ef_wu"] is not None:
            m[2].metric(f"{ef_status(res['ef_wu'])} E-factor (+WU)", res["ef_wu"])
        if res["pmi"] is not None:
            m[3].metric(f"{pmi_status(res['pmi'])} PMI", res["pmi"])
        m2 = st.columns(4)
        if res["ef_plus"] is not None:
            m2[0].metric("E⁺ factor", res["ef_plus"])
        if res["energy"]:
            m2[1].metric("Energy E-factor", f"{res['energy']['ef_energy']} kg CO₂/g",
                         help=f"{res['energy']['kwh']} kWh · CI={CI}")
        if res["ty"]:
            m2[2].metric("Theoretical yield", f"{res['ty']:.4f} g")
        if res["py"] is not None:
            m2[3].metric("% Yield", f"{res['py']} %")

        st.caption(f"PMI computed on **{res['product_basis_label']}** product mass · "
                   f"AE = MW(product)/Σ MW(reagents) · E-factor = waste/product · "
                   f"E⁺ = E-factor(total) + E-factor(energy)")

        st.subheader("Stoichiometry")
        stoich = []
        for r, equiv in res["rows"]:
            stoich.append({
                "Compound": r.name or r.formula or "—",
                "Role": r.role, "CAS": r.cas, "MW": r.mw,
                "Mass (g)": r.mass, "Vol (mL)": r.volume, "mmol": r.mmol,
                "Equiv": round(equiv, 2) if equiv else None, "Purity %": r.purity,
                "Limiting": "★" if r is res["limiting"] else "",
            })
        st.dataframe(pd.DataFrame(stoich), use_container_width=True, hide_index=True)

        if res["biocats"]:
            st.subheader("Biocatalyst summary")
            bio = []
            lm = res["limiting"].mmol
            for x in res["biocats"]:
                amt = x.enzyme_amount
                loading = round(amt * 1000 / lm, 1) if (amt and lm) else None
                total_u = round(amt * 1000 * x.enzyme_activity, 0) if (amt and x.enzyme_activity) else None
                bio.append({
                    "Enzyme": x.enzyme_name, "EC": x.enzyme_code, "Form": x.enzyme_form,
                    "Amount": f"{amt} {'mL' if x.enzyme_form=='liquid' else 'g'}" if amt else "—",
                    "Loading (mg/mmol)": loading, "Total (U)": total_u,
                })
            st.dataframe(pd.DataFrame(bio), use_container_width=True, hide_index=True)

        # Export
        st.divider()
        st.subheader("Export")
        df = build_summary_df(st.session_state.reagents)
        e = st.columns(4)
        e[0].download_button("⬇️ CSV", df.to_csv(index=False).encode(),
                             "reaction_summary.csv", "text/csv", use_container_width=True)
        payload = {"meta": res["meta"], "metrics": {
            k: res[k] for k in ("ae", "ef", "ef_wu", "pmi", "pmi_wu", "ef_plus", "ty", "py")}}
        e[1].download_button("⬇️ JSON", json.dumps(payload, indent=2, default=str).encode(),
                             "reaction_report.json", "application/json", use_container_width=True)

        docx_bytes = build_docx(res["meta"], res, df)
        if docx_bytes:
            e[2].download_button(
                "⬇️ Word", docx_bytes, "reaction_report.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        else:
            e[2].caption("Word export needs `python-docx`.")

        pdf_bytes = build_pdf(res["meta"], res, df)
        if pdf_bytes:
            e[3].download_button("⬇️ PDF", pdf_bytes, "reaction_report.pdf",
                                 "application/pdf", use_container_width=True)
        else:
            e[3].caption("PDF export needs `fpdf2`.")

    else:  # flow
        st.subheader("Flow setup — pumps")
        pump_rows = []
        for i, p in enumerate(res["pumps"]):
            name = p.enzyme_name if (p.role == "catalyst" and p.cat_type == "biological") else p.name
            pump_rows.append({
                "Pump": f"P{str(i+1).zfill(2)}", "Role": p.role,
                "Compound": name or "—", "CAS": p.cas, "MW": p.mw, "mmol": p.mmol,
                "Limiting": "★" if p.is_limiting else "",
            })
        st.dataframe(pd.DataFrame(pump_rows), use_container_width=True, hide_index=True)
        fp = res["flow_product"]
        st.subheader("Product")
        st.write(f"**{fp.name or '—'}** · {fp.formula or '—'} · MW {fp.mw or '—'}")
        if res.get("exp_yield"):
            st.metric("Experimental yield", f"{res['exp_yield']} g")
        st.divider()
        st.subheader("Export")
        df = build_summary_df(res["pumps"] + [fp])
        e = st.columns(4)
        e[0].download_button("⬇️ CSV", df.to_csv(index=False).encode(),
                             "flow_summary.csv", "text/csv", use_container_width=True)
        payload = {"meta": res["meta"], "exp_yield": res.get("exp_yield"),
                   "pumps": [asdict(p) for p in res["pumps"]],
                   "product": asdict(fp)}
        e[1].download_button("⬇️ JSON", json.dumps(payload, indent=2, default=str).encode(),
                             "flow_report.json", "application/json", use_container_width=True)
        docx_bytes = build_docx_flow(res["meta"], res)
        if docx_bytes:
            e[2].download_button(
                "⬇️ Word", docx_bytes, "flow_report.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
        else:
            e[2].caption("Word export needs `python-docx`.")
        e[3].caption("PDF: use Word → Save as PDF.")
